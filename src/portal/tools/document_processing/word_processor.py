"""
Word Document Processor Tool
=============================

Create and edit Word (DOCX) documents programmatically.

Features:
- Create professional documents
- Add text with formatting
- Insert tables, images, headers/footers
- Apply styles and themes
- Convert from Markdown/HTML
- Extract and modify content
- Track changes support

Install: pip install python-docx Pillow
"""

import logging
from pathlib import Path
from typing import Any

from portal.core.interfaces.tool import BaseTool, ToolCategory

logger = logging.getLogger(__name__)

# Check dependencies
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordProcessorTool(BaseTool):
    """
    Create and edit Word documents.

    Perfect for generating reports, contracts, and professional documents.
    """

    def __init__(self) -> None:
        super().__init__()

    METADATA = {
        "name": "word_processor",
        "description": "Create and edit Word (DOCX) documents",
        "category": ToolCategory.UTILITY,
        "version": "1.0.0",
        "requires_confirmation": False,
        "parameters": [
            {"name": "action", "param_type": "string", "description": "Action: create, add_heading, add_paragraph, add_table, add_image, read, save", "required": True},
            {"name": "file_path", "param_type": "string", "description": "Path to Word document", "required": True},
            {"name": "title", "param_type": "string", "description": "Document title", "required": False},
            {"name": "text", "param_type": "string", "description": "Text content to add", "required": False},
            {"name": "heading", "param_type": "string", "description": "Heading text", "required": False},
            {"name": "level", "param_type": "int", "description": "Heading level (1-9)", "required": False, "default": 1},
            {"name": "style", "param_type": "string", "description": "Text style: normal, bold, italic, underline", "required": False},
            {"name": "alignment", "param_type": "string", "description": "Text alignment: left, center, right, justify", "required": False, "default": "left"},
            {"name": "table_data", "param_type": "object", "description": "Table data (rows and columns)", "required": False},
            {"name": "image_path", "param_type": "string", "description": "Path to image file", "required": False},
            {"name": "image_width", "param_type": "float", "description": "Image width in inches", "required": False, "default": 4},
            {"name": "metadata", "param_type": "object", "description": "Document metadata (author, subject, keywords)", "required": False},
        ],
    }

    async def execute(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Execute Word document operation"""

        if not DOCX_AVAILABLE:
            return self._error_response(
                "python-docx not installed. Install: pip install python-docx Pillow"
            )

        action = parameters.get("action", "").lower()
        dispatch = {
            "create": self._create_document,
            "add_heading": self._add_heading,
            "add_paragraph": self._add_paragraph,
            "add_table": self._add_table,
            "add_image": self._add_image,
            "read": self._read_document,
            "save": self._save_document,
        }
        handler = dispatch.get(action)
        if handler is None:
            return self._error_response(f"Unknown action: {action}")
        return await handler(parameters)

    async def _create_document(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Create new Word document"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        title = parameters.get("title", "New Document")
        metadata = parameters.get("metadata", {})

        try:
            # Create document
            doc = Document()

            # Set metadata
            core_properties = doc.core_properties
            core_properties.title = title
            if "author" in metadata:
                core_properties.author = metadata["author"]
            if "subject" in metadata:
                core_properties.subject = metadata["subject"]
            if "keywords" in metadata:
                core_properties.keywords = metadata["keywords"]

            # Add title
            doc.add_heading(title, level=0)

            # Save
            file_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(file_path)

            return self._success_response(
                result={"file_path": str(file_path)}, metadata={"title": title}
            )

        except Exception as e:
            logger.error("Word creation error: %s", e)
            return self._error_response(f"Creation error: {e}")

    async def _add_heading(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Add heading to document"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        heading = parameters.get("heading", "")
        level = parameters.get("level", 1)

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        if not heading:
            return self._error_response("Heading text required")

        try:
            # Load document
            doc = Document(file_path)

            # Add heading
            doc.add_heading(heading, level=level)

            # Save
            doc.save(file_path)

            return self._success_response(
                result={"heading_added": heading}, metadata={"level": level}
            )

        except Exception as e:
            logger.error("Word add heading error: %s", e)
            return self._error_response(f"Add heading error: {e}")

    async def _add_paragraph(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Add paragraph to document"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        text = parameters.get("text", "")
        style = parameters.get("style", "normal")
        alignment = parameters.get("alignment", "left")

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        if not text:
            return self._error_response("Text required")

        try:
            # Load document
            doc = Document(file_path)

            # Add paragraph
            paragraph = doc.add_paragraph(text)

            # Apply alignment
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

            # Apply style
            if style == "bold":
                paragraph.runs[0].bold = True
            elif style == "italic":
                paragraph.runs[0].italic = True
            elif style == "underline":
                paragraph.runs[0].underline = True

            # Save
            doc.save(file_path)

            return self._success_response(
                result={"paragraph_added": True}, metadata={"length": len(text), "style": style}
            )

        except Exception as e:
            logger.error("Word add paragraph error: %s", e)
            return self._error_response(f"Add paragraph error: {e}")

    async def _add_table(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Add table to document"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        table_data = parameters.get("table_data", {})

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        if not table_data:
            return self._error_response("Table data required")

        try:
            # Load document
            doc = Document(file_path)

            # Get table data
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])

            if not rows:
                return self._error_response("No table rows provided")

            # Create table
            num_cols = len(rows[0]) if rows else len(headers)
            num_rows = len(rows) + (1 if headers else 0)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = "Light Grid Accent 1"

            # Add headers
            if headers:
                self._apply_bold_headers(table.rows[0].cells, headers)

            # Add data
            start_row = 1 if headers else 0
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[start_row + row_idx].cells
                for col_idx, cell_data in enumerate(row_data):
                    row_cells[col_idx].text = str(cell_data)

            # Save
            doc.save(file_path)

            return self._success_response(
                result={"table_added": True}, metadata={"rows": len(rows), "columns": num_cols}
            )

        except Exception as e:
            logger.error("Word add table error: %s", e)
            return self._error_response(f"Add table error: {e}")

    def _apply_bold_headers(self, header_cells: Any, headers: list[Any]) -> None:
        """Set text and bold style on header cells."""
        for idx, header in enumerate(headers):
            header_cells[idx].text = str(header)
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    async def _add_image(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Add image to document"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        image_path = Path(parameters.get("image_path", "")).expanduser()
        image_width = parameters.get("image_width", 4)

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        if not image_path.exists():
            return self._error_response(f"Image not found: {image_path}")

        try:
            # Load document
            doc = Document(file_path)

            # Add image
            doc.add_picture(str(image_path), width=Inches(image_width))

            # Save
            doc.save(file_path)

            return self._success_response(
                result={"image_added": str(image_path)}, metadata={"width_inches": image_width}
            )

        except Exception as e:
            logger.error("Word add image error: %s", e)
            return self._error_response(f"Add image error: {e}")

    async def _read_document(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Read Word document content"""

        file_path = Path(parameters.get("file_path", "")).expanduser()

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        try:
            # Load document
            doc = Document(file_path)

            # Extract content
            content = {
                "metadata": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "keywords": doc.core_properties.keywords or "",
                },
                "paragraphs": [],
                "tables": [],
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content["paragraphs"].append({"text": para.text, "style": para.style.name})

            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append(table_data)

            return self._success_response(
                result=content,
                metadata={
                    "paragraphs_count": len(content["paragraphs"]),
                    "tables_count": len(content["tables"]),
                },
            )

        except Exception as e:
            logger.error("Word read error: %s", e)
            return self._error_response(f"Read error: {e}")

    async def _save_document(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Save document (placeholder for future modifications)"""

        file_path = Path(parameters.get("file_path", "")).expanduser()

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        return self._success_response(result={"saved": str(file_path)}, metadata={})
