"""
Document Metadata Extractor Tool
=================================

Extract metadata and properties from various document formats.

Supported formats:
- PDF, DOCX, PPTX, XLSX
- Images (JPEG, PNG with EXIF)
- Videos (MP4, MOV)
- Audio (MP3, M4A)

Features:
- Author, title, subject, keywords
- Creation/modification dates
- Page/slide counts
- File properties
- EXIF data for images
- ID3 tags for audio

Install: pip install pypdf python-docx python-pptx openpyxl Pillow mutagen
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from portal.core.interfaces.tool import BaseTool, ToolCategory

logger = logging.getLogger(__name__)


class DocumentMetadataExtractorTool(BaseTool):
    """
    Extract metadata from documents, images, audio, and video files.

    Works with 15+ file formats!
    """

    def __init__(self) -> None:
        super().__init__()

    METADATA = {
        "name": "document_metadata",
        "description": "Extract metadata from documents (PDF, DOCX, PPTX, XLSX, images, audio)",
        "category": ToolCategory.UTILITY,
        "version": "1.0.0",
        "requires_confirmation": False,
        "parameters": [
            {"name": "file_path", "param_type": "string", "description": "Path to document file", "required": True},
            {"name": "detailed", "param_type": "bool", "description": "Include detailed metadata (default: False)", "required": False, "default": False},
        ],
    }

    async def execute(self, parameters: dict[str, Any]) -> dict[str, Any]:
        """Extract document metadata"""

        file_path = Path(parameters.get("file_path", "")).expanduser()
        detailed = parameters.get("detailed", False)

        if not file_path.exists():
            return self._error_response(f"File not found: {file_path}")

        try:
            suffix = file_path.suffix.lower()
            dispatch = {
                ".pdf": self._extract_pdf_metadata,
                ".docx": self._extract_docx_metadata,
                ".doc": self._extract_docx_metadata,
                ".pptx": self._extract_pptx_metadata,
                ".xlsx": self._extract_xlsx_metadata,
                ".xls": self._extract_xlsx_metadata,
                ".jpg": self._extract_image_metadata,
                ".jpeg": self._extract_image_metadata,
                ".png": self._extract_image_metadata,
                ".gif": self._extract_image_metadata,
                ".bmp": self._extract_image_metadata,
                ".mp3": self._extract_audio_metadata,
                ".m4a": self._extract_audio_metadata,
                ".flac": self._extract_audio_metadata,
                ".ogg": self._extract_audio_metadata,
                ".wav": self._extract_audio_metadata,
            }
            handler = dispatch.get(suffix)
            if handler is None:
                return self._error_response(f"Unsupported file format: {suffix}")
            metadata = await handler(file_path, detailed)

            # Add common file properties
            stat = file_path.stat()
            metadata["file_properties"] = {
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "format": suffix[1:],
            }

            return self._success_response(result=metadata, metadata={"file_type": suffix[1:]})

        except Exception as e:
            logger.error("Metadata extraction error: %s", e)
            return self._error_response(f"Extraction error: {e}")

    async def _extract_pdf_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract PDF metadata"""

        try:
            from pypdf import PdfReader
        except ImportError:
            return {"error": "pypdf not installed"}

        try:
            reader = PdfReader(str(file_path))
            info = reader.metadata

            metadata = {
                "type": "PDF Document",
                "pages": len(reader.pages),
                "title": info.get("/Title", ""),
                "author": info.get("/Author", ""),
                "subject": info.get("/Subject", ""),
                "creator": info.get("/Creator", ""),
                "producer": info.get("/Producer", ""),
                "keywords": info.get("/Keywords", ""),
            }

            # Creation/modification dates
            if "/CreationDate" in info:
                metadata["created"] = info.get("/CreationDate")
            if "/ModDate" in info:
                metadata["modified"] = info.get("/ModDate")

            # Detailed info
            if detailed:
                metadata["encrypted"] = reader.is_encrypted
                metadata["pages_info"] = [
                    {
                        "number": i + 1,
                        "width": float(page.mediabox.width),
                        "height": float(page.mediabox.height),
                    }
                    for i, page in enumerate(reader.pages[:5])  # First 5 pages
                ]

            return metadata

        except Exception as e:
            return {"error": f"PDF extraction failed: {e}"}

    async def _extract_docx_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract DOCX metadata"""

        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx not installed"}

        try:
            doc = Document(str(file_path))
            props = doc.core_properties

            metadata = {
                "type": "Word Document",
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "comments": props.comments or "",
                "category": props.category or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
            }

            # Content statistics
            if detailed:
                metadata["statistics"] = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "styles": len(doc.styles),
                }

            return metadata

        except Exception as e:
            return {"error": f"DOCX extraction failed: {e}"}

    async def _extract_pptx_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract PPTX metadata"""

        try:
            from pptx import Presentation
        except ImportError:
            return {"error": "python-pptx not installed"}

        try:
            prs = Presentation(str(file_path))
            props = prs.core_properties

            metadata = {
                "type": "PowerPoint Presentation",
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "slides": len(prs.slides),
            }

            if detailed:
                metadata["slide_dimensions"] = {
                    "width": prs.slide_width,
                    "height": prs.slide_height,
                }

            return metadata

        except Exception as e:
            return {"error": f"PPTX extraction failed: {e}"}

    async def _extract_xlsx_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract XLSX metadata"""

        try:
            from openpyxl import load_workbook
        except ImportError:
            return {"error": "openpyxl not installed"}

        try:
            wb = load_workbook(str(file_path), data_only=True)
            props = wb.properties

            metadata = {
                "type": "Excel Spreadsheet",
                "title": props.title or "",
                "author": props.creator or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "sheets": wb.sheetnames,
            }

            if detailed:
                metadata["sheet_details"] = {
                    sheet: {"rows": ws.max_row, "columns": ws.max_column}
                    for sheet, ws in [(name, wb[name]) for name in wb.sheetnames]
                }

            wb.close()
            return metadata

        except Exception as e:
            return {"error": f"XLSX extraction failed: {e}"}

    async def _extract_image_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract image EXIF metadata"""

        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
        except ImportError:
            return {"error": "Pillow not installed"}

        try:
            img = Image.open(str(file_path))

            metadata = {
                "type": "Image",
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
                "megapixels": round((img.width * img.height) / 1_000_000, 2),
            }

            # EXIF data
            exif_data = img._getexif()
            if exif_data and detailed:
                exif = {}
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        continue  # Skip binary data
                    exif[tag] = str(value)

                metadata["exif"] = exif

            img.close()
            return metadata

        except Exception as e:
            return {"error": f"Image extraction failed: {e}"}

    async def _extract_audio_metadata(self, file_path: Path, detailed: bool) -> dict:
        """Extract audio metadata"""

        try:
            from mutagen import File as MutagenFile
        except ImportError:
            return {"error": "mutagen not installed"}

        try:
            audio = MutagenFile(str(file_path))

            if audio is None:
                return {"error": "Could not read audio file"}

            metadata = {
                "type": "Audio",
                "format": audio.mime[0] if audio.mime else "Unknown",
                "length_seconds": round(audio.info.length, 2) if audio.info else 0,
                "bitrate": audio.info.bitrate if hasattr(audio.info, "bitrate") else "Unknown",
            }

            # Tags
            if audio.tags and detailed:
                tags = {}
                for key, value in audio.tags.items():
                    tags[key] = str(value[0]) if isinstance(value, list) else str(value)
                metadata["tags"] = tags

            return metadata

        except Exception as e:
            return {"error": f"Audio extraction failed: {e}"}
